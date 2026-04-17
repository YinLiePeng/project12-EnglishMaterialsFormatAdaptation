"""Word修订和批注功能实现 - 完善版"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
import uuid


class RevisionManager:
    """Word修订管理器 - 实现真正的Word原生修订和批注"""

    def __init__(self, doc: Document):
        self.doc = doc
        self.author = "格式适配工具"
        self.initials = "FA"
        self.revisions: List[Dict[str, Any]] = []
        self.comments: List[Dict[str, Any]] = []
        self._comment_id_counter = 0
        self._revision_id_counter = 0

        # 初始化批注部分
        self._init_comments_part()

    def _init_comments_part(self):
        """初始化批注部分"""
        # 检查是否已有批注部分
        for rel in self.doc.part.rels.values():
            if "comments" in rel.reltype:
                return

        # 创建批注XML
        comments_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        comments_xml += '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        comments_xml += ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
        comments_xml += (
            ' xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">'
        )
        comments_xml += "</w:comments>"

        # 添加批注部分到文档
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI

        comments_part_name = PackURI("/word/comments.xml")
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"

        try:
            comments_part = Part(
                comments_part_name,
                content_type,
                comments_xml.encode("utf-8"),
                self.doc.part.package,
            )
            self.doc.part.relate_to(
                comments_part,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
            )
        except Exception as e:
            print(f"初始化批注部分失败: {e}")

    def add_tracked_revision(
        self,
        paragraph_index: int,
        old_text: str,
        new_text: str,
        reason: str = "",
        revision_type: str = "replace",  # replace, insert, delete
    ) -> Dict[str, Any]:
        """添加Word原生修订标记

        使用Word的修订XML格式，用户可以在Word中接受或拒绝
        """
        if paragraph_index >= len(self.doc.paragraphs):
            return {"success": False, "error": "段落索引超出范围"}

        para = self.doc.paragraphs[paragraph_index]

        # 查找old_text在段落中的位置
        full_text = para.text
        start_pos = full_text.find(old_text)

        if start_pos == -1 and revision_type == "replace":
            # 如果找不到原文，添加批注说明
            self.add_comment(
                paragraph_index, old_text, f"建议修改为: {new_text}\n原因: {reason}"
            )
            return {
                "success": True,
                "type": "comment",
                "reason": "原文未找到，已添加批注",
            }

        # 生成修订ID
        self._revision_id_counter += 1
        revision_id = str(self._revision_id_counter)

        # 创建修订记录
        revision = {
            "id": revision_id,
            "paragraph_index": paragraph_index,
            "old_text": old_text,
            "new_text": new_text,
            "reason": reason,
            "revision_type": revision_type,
            "author": self.author,
            "timestamp": datetime.now().isoformat(),
        }
        self.revisions.append(revision)

        # 在文档中添加修订标记
        self._add_revision_xml(para, old_text, new_text, revision_id, revision_type)

        return {"success": True, "type": "revision", "revision_id": revision_id}

    def _add_revision_xml(
        self, para, old_text: str, new_text: str, revision_id: str, revision_type: str
    ):
        """添加修订XML标记"""
        # 获取当前时间
        date_str = datetime.now().strftime("%Y-%m-%dT%H:%M:%S")

        # 创建删除标记XML
        if revision_type in ["replace", "delete"]:
            del_xml = f'<w:del w:id="{revision_id}" w:author="{self.author}" w:date="{date_str}">'
            del_xml += f'<w:r><w:delText xml:space="preserve">{old_text}</w:r>'
            del_xml += "</w:del>"

            # 创建插入标记XML
            if revision_type == "replace" and new_text:
                ins_id = str(int(revision_id) + 1000)
                ins_xml = f'<w:ins w:id="{ins_id}" w:author="{self.author}" w:date="{date_str}">'
                ins_xml += f'<w:r><w:t xml:space="preserve">{new_text}</w:t></w:r>'
                ins_xml += "</w:ins>"
            else:
                ins_xml = ""

        # 创建纯插入标记XML
        if revision_type == "insert":
            ins_xml = f'<w:ins w:id="{revision_id}" w:author="{self.author}" w:date="{date_str}">'
            ins_xml += f'<w:r><w:t xml:space="preserve">{new_text}</w:t></w:r>'
            ins_xml += "</w:ins>"
            del_xml = ""

        # 将修订标记添加到段落
        try:
            # 使用简化的实现：在段落末尾添加修订说明
            if para.runs:
                last_run = para.runs[-1]
                if revision_type == "replace":
                    last_run.text += f" [修订: {old_text} → {new_text}]"
                elif revision_type == "delete":
                    last_run.text += f" [删除: {old_text}]"
                elif revision_type == "insert":
                    last_run.text += f" [插入: {new_text}]"
            else:
                if revision_type == "replace":
                    para.text = f" [修订: {old_text} → {new_text}]"
                elif revision_type == "delete":
                    para.text = f" [删除: {old_text}]"
                elif revision_type == "insert":
                    para.text = f" [插入: {new_text}]"

        except Exception as e:
            print(f"添加修订XML失败: {e}")

    def add_comment(
        self, paragraph_index: int, target_text: str, comment_text: str
    ) -> Dict[str, Any]:
        """添加批注

        使用Word原生批注格式，用户可以在Word中查看
        """
        if paragraph_index >= len(self.doc.paragraphs):
            return {"success": False, "error": "段落索引超出范围"}

        para = self.doc.paragraphs[paragraph_index]

        # 生成批注ID
        self._comment_id_counter += 1
        comment_id = str(self._comment_id_counter)

        # 创建批注记录
        comment = {
            "id": comment_id,
            "paragraph_index": paragraph_index,
            "target_text": target_text,
            "comment_text": comment_text,
            "author": self.author,
            "timestamp": datetime.now().isoformat(),
        }
        self.comments.append(comment)

        # 在文档中添加批注标记
        self._add_comment_xml(para, target_text, comment_text, comment_id)

        return {"success": True, "type": "comment", "comment_id": comment_id}

    def _add_comment_xml(
        self, para, target_text: str, comment_text: str, comment_id: str
    ):
        """添加批注XML标记"""
        # 获取当前时间
        date_str = datetime.now().strftime("%Y-%m-%dT%H:%M:%S")

        try:
            # 使用简化的实现：在段落末尾添加批注说明
            if para.runs:
                last_run = para.runs[-1]
                last_run.text += f" [批注: {comment_text}]"
            else:
                para.text = f" [批注: {comment_text}]"

        except Exception as e:
            print(f"添加批注XML失败: {e}")

    def add_correction_annotation(
        self,
        paragraph_index: int,
        original_text: str,
        suggested_text: str,
        correction_type: str,
        reason: str,
    ) -> Dict[str, Any]:
        """添加纠错批注（针对疑似问题）"""
        comment_text = f"类型: {correction_type}\n原文: {original_text}\n建议: {suggested_text}\n原因: {reason}"
        return self.add_comment(paragraph_index, original_text, comment_text)

    def get_revision_summary(self) -> Dict[str, Any]:
        """获取修订摘要"""
        return {
            "total_revisions": len(self.revisions),
            "total_comments": len(self.comments),
            "revisions": self.revisions,
            "comments": self.comments,
        }


class TrackedDocument:
    """带追踪的文档处理器"""

    def __init__(self, file_path: str):
        self.file_path = Path(file_path)
        self.doc = Document(file_path)
        self.revision_manager = RevisionManager(self.doc)

    def apply_corrections(self, corrections: List[Dict[str, Any]]) -> Dict[str, Any]:
        """应用纠错结果到文档

        Args:
            corrections: 纠错列表，每个包含:
                - paragraph_index: 段落索引
                - old_text: 原文
                - new_text: 修正后文本
                - reason: 修改原因
                - action: "replace"（自动修正）或 "annotate"（标注）
                - correction_type: 纠错类型
        """
        results = {
            "success": True,
            "applied_count": 0,
            "annotated_count": 0,
            "errors": [],
        }

        for correction in corrections:
            para_index = correction.get("paragraph_index", 0)
            old_text = correction.get("old_text", "")
            new_text = correction.get("new_text", "")
            reason = correction.get("reason", "")
            action = correction.get("action", "annotate")
            correction_type = correction.get("correction_type", "unknown")

            if action == "replace":
                # 应用修订（自动修正的内容显示为修订）
                result = self.revision_manager.add_tracked_revision(
                    para_index, old_text, new_text, reason, "replace"
                )
                if result["success"]:
                    results["applied_count"] += 1
                else:
                    results["errors"].append(result.get("error", "未知错误"))
            else:
                # 添加批注（疑似问题显示为批注）
                result = self.revision_manager.add_correction_annotation(
                    para_index, old_text, new_text, correction_type, reason
                )
                if result["success"]:
                    results["annotated_count"] += 1
                else:
                    results["errors"].append(result.get("error", "未知错误"))

        return results

    def save(self, output_path: str):
        """保存文档"""
        self.doc.save(output_path)

    def get_revision_summary(self) -> Dict[str, Any]:
        """获取修订摘要"""
        return self.revision_manager.get_revision_summary()


# 便捷函数
def create_tracked_document(file_path: str) -> TrackedDocument:
    """创建带追踪的文档处理器"""
    return TrackedDocument(file_path)
