"""PDF转DOCX转换器 - 基于opendataloader_pdf JSON输出"""

from pathlib import Path
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .parser import PDFParser, FontMapper, ColorConverter, AlignmentDetector


class PDFStyleMapper:
    """PDF样式映射器 - 将PDF样式映射为DOCX样式"""
    
    @classmethod
    def map_font(cls, pdf_font: str) -> str:
        """映射PDF字体到DOCX字体"""
        return FontMapper.map_font(pdf_font)
    
    @classmethod
    def map_color(cls, color_str: str) -> RGBColor:
        """映射PDF颜色到DOCX颜色"""
        return ColorConverter.parse_rgb(color_str)
    
    @classmethod
    def map_paragraph_style(cls, element: Dict[str, Any]) -> Dict[str, Any]:
        """映射段落样式"""
        bbox = element.get('bounding box', [0, 0, 0, 0])
        alignment = AlignmentDetector.detect_from_bbox(bbox)
        
        return {
            "font": {
                "name": cls.map_font(element.get('font', '')),
                "size": Pt(element.get('font size', 12)),
                "bold": FontMapper.is_bold(element.get('font', '')),
                "italic": FontMapper.is_italic(element.get('font', '')),
                "color": cls.map_color(element.get('text color', '')),
            },
            "format": {
                "alignment": AlignmentDetector.get_wd_align(alignment),
                "line_spacing": 1.25,
            },
        }


class PDFToDocxConverter:
    """PDF转DOCX转换器 - 直接基于JSON生成高质量DOCX"""
    
    def __init__(self):
        self.style_mapper = PDFStyleMapper()
    
    def convert(
        self,
        pdf_path: str,
        output_path: str = None,
        include_tables: bool = True,
        include_images: bool = True,
        preserve_format: bool = True,
    ) -> str:
        """将PDF转换为DOCX
        
        Args:
            pdf_path: PDF文件路径
            output_path: 输出DOCX路径（可选）
            include_tables: 是否包含表格
            include_images: 是否包含图片
            preserve_format: 是否保留原格式
        
        Returns:
            输出DOCX文件路径
        """
        pdf_path = Path(pdf_path)
        if not output_path:
            output_path = str(pdf_path.with_suffix('.docx'))
        
        # 确保JSON已生成
        json_path = pdf_path.with_suffix('.json')
        images_dir = Path(str(pdf_path.with_suffix('')) + '_images')
        
        if not json_path.exists():
            import opendataloader_pdf
            opendataloader_pdf.convert(
                input_path=[str(pdf_path)],
                output_dir=str(pdf_path.parent),
                format="json"
            )
        
        # 加载JSON
        import json
        with open(json_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        
        # 创建DOCX文档
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 获取页面尺寸信息
        page_width = 595.0  # 默认A4宽度（points）
        page_height = 842.0  # 默认A4高度（points）
        
        # 获取所有元素并按页面和Y坐标排序
        all_elements = json_data.get('kids', [])
        
        # 过滤掉header/footer中的元素，避免重复
        filtered_elements = self._filter_elements(all_elements)
        
        # 按页面和Y坐标排序（PDF坐标系Y从底部开始）
        filtered_elements.sort(key=lambda e: (
            e.get('page number', 0),
            -(e.get('bounding box', [0, 0, 0, 0])[1] if len(e.get('bounding box', [])) >= 2 else 0)
        ))
        
        # 处理元素
        current_page = 0
        prev_element = None
        
        for i, element in enumerate(filtered_elements):
            page_num = element.get('page number', 0)
            
            # 分页控制
            if page_num > current_page:
                doc.add_page_break()
                current_page = page_num
                prev_element = None
            
            # 计算段落间距
            space_before = None
            if prev_element and preserve_format:
                space_before = self._calculate_space_before(prev_element, element, page_height)
            
            elem_type = element.get('type')
            
            if elem_type == 'heading':
                self._add_heading(doc, element, preserve_format, space_before)
            elif elem_type == 'paragraph':
                self._add_paragraph(doc, element, preserve_format, space_before, page_width)
            elif elem_type == 'table' and include_tables:
                self._add_table(doc, element, preserve_format, space_before)
            elif elem_type == 'image' and include_images:
                self._add_image(doc, element, pdf_path, space_before)
            elif elem_type == 'list':
                self._add_list(doc, element, preserve_format, space_before)
            
            prev_element = element
        
        # 保存文档
        doc.save(output_path)
        return output_path
    
    def _filter_elements(self, all_elements: List[Dict]) -> List[Dict]:
        """过滤元素，移除header/footer中的重复元素"""
        header_footer_ids = set()
        
        for elem in all_elements:
            if elem.get('type') in ['header', 'footer']:
                for kid in elem.get('kids', []):
                    header_footer_ids.add(kid.get('id'))
        
        filtered = []
        for elem in all_elements:
            if elem.get('id') in header_footer_ids:
                continue
            if elem.get('type') in ['header', 'footer']:
                # 处理页眉页脚
                continue
            filtered.append(elem)
        
        return filtered
    
    def _calculate_space_before(self, prev_element: Dict, current_element: Dict, page_height: float) -> Optional[Pt]:
        """计算段前间距"""
        try:
            prev_bbox = prev_element.get('bounding box', [0, 0, 0, 0])
            curr_bbox = current_element.get('bounding box', [0, 0, 0, 0])
            
            if len(prev_bbox) >= 4 and len(curr_bbox) >= 4:
                # PDF坐标系Y从底部开始
                prev_top = prev_bbox[3]
                curr_bottom = curr_bbox[1]
                
                # 计算垂直间距（points）
                gap = prev_top - curr_bottom
                
                # 如果间距大于一定阈值，添加段前间距
                if gap > 5:  # 大于5 points
                    return Pt(gap)
        except Exception:
            pass
        
        return None
    
    def _calculate_indent(self, bbox: List[float], page_width: float) -> Optional[Pt]:
        """计算首行缩进或左缩进"""
        if not bbox or len(bbox) < 4:
            return None
        
        left_pos = bbox[0]
        
        # 如果左边距大于页面宽度的10%，认为是缩进
        if left_pos > page_width * 0.1:
            return Pt(left_pos)
        
        return None
    
    def _add_heading(self, doc: Document, element: Dict, preserve_format: bool, space_before: Optional[Pt] = None):
        """添加标题"""
        content = element.get('content', '')
        if not content:
            return
        
        heading_level = element.get('heading level', 1)
        level = min(max(heading_level, 1), 4)  # 限制在1-4级
        
        heading = doc.add_heading(content, level=level)
        
        if preserve_format:
            # 应用字体样式
            for run in heading.runs:
                run.font.name = FontMapper.map_font(element.get('font', ''))
                run.font.size = Pt(element.get('font size', 12))
                run.font.bold = FontMapper.is_bold(element.get('font', ''))
                run.font.italic = FontMapper.is_italic(element.get('font', ''))
                run.font.color.rgb = ColorConverter.parse_rgb(element.get('text color', ''))
                run._element.rPr.rFonts.set(qn('w:eastAsia'), FontMapper.map_font(element.get('font', '')))
            
            # 应用对齐方式
            bbox = element.get('bounding box', [0, 0, 0, 0])
            alignment = AlignmentDetector.detect_from_bbox(bbox)
            heading.alignment = AlignmentDetector.get_wd_align(alignment)
            
            # 应用段前间距
            if space_before:
                heading.paragraph_format.space_before = space_before
    
    def _add_paragraph(self, doc: Document, element: Dict, preserve_format: bool, 
                       space_before: Optional[Pt] = None, page_width: float = 595.0):
        """添加段落"""
        content = element.get('content', '')
        if not content:
            return
        
        para = doc.add_paragraph(content)
        
        if preserve_format:
            # 应用字体样式
            for run in para.runs:
                run.font.name = FontMapper.map_font(element.get('font', ''))
                run.font.size = Pt(element.get('font size', 12))
                run.font.bold = FontMapper.is_bold(element.get('font', ''))
                run.font.italic = FontMapper.is_italic(element.get('font', ''))
                run.font.color.rgb = ColorConverter.parse_rgb(element.get('text color', ''))
                run._element.rPr.rFonts.set(qn('w:eastAsia'), FontMapper.map_font(element.get('font', '')))
            
            # 应用对齐方式
            bbox = element.get('bounding box', [0, 0, 0, 0])
            alignment = AlignmentDetector.detect_from_bbox(bbox)
            para.alignment = AlignmentDetector.get_wd_align(alignment)
            
            # 应用缩进
            indent = self._calculate_indent(bbox, page_width)
            if indent:
                para.paragraph_format.left_indent = indent
            
            # 应用段前间距
            if space_before:
                para.paragraph_format.space_before = space_before
    
    def _add_table(self, doc: Document, element: Dict, preserve_format: bool, 
                   space_before: Optional[Pt] = None):
        """添加表格"""
        rows = element.get('rows', [])
        if not rows:
            return
        
        num_rows = len(rows)
        num_cols = max(len(row.get('cells', [])) for row in rows) if rows else 0
        
        if num_rows == 0 or num_cols == 0:
            return
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        # 应用段前间距
        if space_before:
            table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = space_before
        
        # 填充单元格内容
        for row_idx, row in enumerate(rows):
            cells = row.get('cells', [])
            for col_idx, cell in enumerate(cells):
                if col_idx >= num_cols:
                    break
                
                cell_text = self._extract_cell_text(cell)
                table.cell(row_idx, col_idx).text = cell_text
                
                # 处理合并单元格
                row_span = cell.get('row span', 1)
                col_span = cell.get('column span', 1)
                
                if row_span > 1 or col_span > 1:
                    # 合并单元格
                    start_row = row_idx
                    end_row = row_idx + row_span - 1
                    start_col = col_idx
                    end_col = col_idx + col_span - 1
                    
                    # 确保不超出表格范围
                    end_row = min(end_row, num_rows - 1)
                    end_col = min(end_col, num_cols - 1)
                    
                    if start_row != end_row or start_col != end_col:
                        try:
                            start_cell = table.cell(start_row, start_col)
                            end_cell = table.cell(end_row, end_col)
                            start_cell.merge(end_cell)
                        except Exception:
                            pass
        
        # 应用表格格式
        if preserve_format:
            # 应用单元格内字体样式
            for row_idx, row in enumerate(rows):
                cells = row.get('cells', [])
                for col_idx, cell in enumerate(cells):
                    if col_idx >= num_cols:
                        break
                    
                    # 获取单元格内第一个段落的字体信息
                    cell_kids = cell.get('kids', [])
                    if cell_kids:
                        first_kid = cell_kids[0]
                        if first_kid.get('type') in ['paragraph', 'heading']:
                            try:
                                cell_obj = table.cell(row_idx, col_idx)
                                for para in cell_obj.paragraphs:
                                    for run in para.runs:
                                        run.font.name = FontMapper.map_font(first_kid.get('font', ''))
                                        run.font.size = Pt(first_kid.get('font size', 12))
                                        run.font.bold = FontMapper.is_bold(first_kid.get('font', ''))
                                        run.font.italic = FontMapper.is_italic(first_kid.get('font', ''))
                                        run.font.color.rgb = ColorConverter.parse_rgb(first_kid.get('text color', ''))
                                        run._element.rPr.rFonts.set(qn('w:eastAsia'), FontMapper.map_font(first_kid.get('font', '')))
                            except Exception:
                                pass
    
    def _extract_cell_text(self, cell: Dict) -> str:
        """提取单元格文本"""
        texts = []
        for kid in cell.get('kids', []):
            if kid.get('type') in ['paragraph', 'heading']:
                texts.append(kid.get('content', ''))
            elif kid.get('type') == 'list':
                for item in kid.get('list items', []):
                    texts.append(item.get('content', ''))
        return '\n'.join(texts)
    
    def _add_image(self, doc: Document, element: Dict, pdf_path: Path, 
                   space_before: Optional[Pt] = None):
        """添加图片"""
        source = element.get('source', '')
        if not source:
            return
        
        # 构建图片路径
        images_dir = Path(str(pdf_path.with_suffix('')) + '_images')
        img_path = images_dir / Path(source).name
        
        if not img_path.exists():
            return
        
        # 检查文件大小，跳过装饰性小图片
        file_size = img_path.stat().st_size
        if file_size < 200:  # 小于200字节的可能是装饰
            return
        
        try:
            # 计算图片宽度（从bounding box）
            bbox = element.get('bounding box', [0, 0, 0, 0])
            if len(bbox) >= 4:
                width_pt = bbox[2] - bbox[0]
                width_inches = width_pt / 72.0  # 1 inch = 72 points
                # 限制最大宽度为页面宽度（约6.5英寸）
                width_inches = min(width_inches, 6.5)
            else:
                width_inches = None
            
            # 添加图片
            if width_inches and width_inches > 0.5:  # 至少0.5英寸才添加
                pic = doc.add_picture(str(img_path), width=Inches(width_inches))
            else:
                pic = doc.add_picture(str(img_path))
            
            # 应用段前间距
            if space_before:
                # 图片段落的段前间距
                pic_para = doc.paragraphs[-1]
                pic_para.paragraph_format.space_before = space_before
                
        except Exception:
            # 图片添加失败则跳过
            pass
    
    def _add_list(self, doc: Document, element: Dict, preserve_format: bool, 
                  space_before: Optional[Pt] = None):
        """添加列表"""
        list_items = element.get('list items', [])
        numbering_style = element.get('numbering style', 'unordered')
        
        for i, item in enumerate(list_items):
            content = item.get('content', '')
            if not content:
                continue
            
            # 构建列表项文本
            if numbering_style == 'arabic numbers' and not content[0].isdigit():
                content = f"{i + 1}. {content}"
            elif numbering_style == 'unordered' and not content.startswith('-'):
                content = f"• {content}"
            
            para = doc.add_paragraph(content, style='List Bullet' if numbering_style == 'unordered' else 'List Number')
            
            if preserve_format:
                # 应用字体样式
                for run in para.runs:
                    run.font.name = FontMapper.map_font(item.get('font', ''))
                    run.font.size = Pt(item.get('font size', 12))
                    run.font.bold = FontMapper.is_bold(item.get('font', ''))
                    run.font.italic = FontMapper.is_italic(item.get('font', ''))
                    run.font.color.rgb = ColorConverter.parse_rgb(item.get('text color', ''))
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), FontMapper.map_font(item.get('font', '')))
            
            # 应用段前间距（仅第一个列表项）
            if i == 0 and space_before:
                para.paragraph_format.space_before = space_before
            
            # 处理嵌套元素
            for kid in item.get('kids', []):
                if kid.get('type') == 'paragraph':
                    self._add_paragraph(doc, kid, preserve_format)
                elif kid.get('type') == 'list':
                    self._add_list(doc, kid, preserve_format)


def convert_pdf_to_docx(
    pdf_path: str,
    output_path: str = None,
    **kwargs
) -> str:
    """便捷函数：将PDF转换为DOCX"""
    converter = PDFToDocxConverter()
    return converter.convert(pdf_path, output_path, **kwargs)
