"""Phase 5 测试脚本 - 最终批量测试和性能评估

测试内容：
1. 使用新的后端流程处理所有19个PDF
2. 测试"保留原格式"模式下的排版效果
3. 评估处理速度
4. 生成最终对比报告
"""

import sys
import time
import json
from pathlib import Path
from typing import Dict, List
from difflib import SequenceMatcher

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from app.services.pdf.converter import PDFToDocxConverter


# 路径配置
PDF_DIR = Path("/mnt/e/Opencode/project12-EnglishMaterialsFormatAdaptation/测试用例/原生PDF")
REF_DIR = Path("/mnt/e/Opencode/project12-EnglishMaterialsFormatAdaptation/测试用例/网上下载的英语资料样例")
OUTPUT_DIR = Path("/mnt/e/Opencode/project12-EnglishMaterialsFormatAdaptation/backend/data/phase5_final_output")


def get_pdf_files() -> List[Path]:
    """获取所有PDF文件"""
    return sorted([f for f in PDF_DIR.glob("*.pdf") if f.is_file()])


def find_reference_docx(pdf_file: Path) -> Path:
    """查找对应的标准DOCX文件"""
    base_name = pdf_file.stem
    ref_file = REF_DIR / f"{base_name}.docx"
    if ref_file.exists():
        return ref_file
    
    for ref in REF_DIR.glob("*.docx"):
        if base_name[:30] in ref.stem or ref.stem[:30] in base_name:
            return ref
    
    return None


def extract_docx_text(docx_path: Path) -> str:
    """提取DOCX中的所有文本"""
    try:
        doc = Document(str(docx_path))
        texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texts.append(cell.text.strip())
        return '\n'.join(texts)
    except Exception:
        return ""


def extract_docx_structure(docx_path: Path) -> Dict:
    """提取DOCX的结构信息"""
    try:
        doc = Document(str(docx_path))
        
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    'text': para.text.strip(),
                    'style': para.style.name if para.style else 'Normal',
                })
        
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables.append(table_data)
        
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_count += 1
        
        return {
            'paragraph_count': len(paragraphs),
            'table_count': len(tables),
            'image_count': image_count,
        }
    except Exception:
        return {'paragraph_count': 0, 'table_count': 0, 'image_count': 0}


def calculate_similarity(text1: str, text2: str) -> float:
    """计算文本相似度"""
    if not text1 and not text2:
        return 1.0
    if not text1 or not text2:
        return 0.0
    return SequenceMatcher(None, text1, text2).ratio()


def test_single_pdf(pdf_file: Path, ref_docx: Path) -> Dict:
    """测试单个PDF"""
    result = {
        'pdf_name': pdf_file.name,
        'ref_name': ref_docx.name if ref_docx else None,
        'success': False,
        'error': None,
        'processing_time': 0,
        'text_similarity': 0.0,
        'structure_similarity': 0.0,
        'generated_paragraphs': 0,
        'reference_paragraphs': 0,
        'generated_tables': 0,
        'reference_tables': 0,
        'generated_images': 0,
        'reference_images': 0,
    }
    
    if not ref_docx:
        result['error'] = "未找到对应的标准DOCX文件"
        return result
    
    try:
        # 生成DOCX（保留原格式模式）
        start_time = time.time()
        converter = PDFToDocxConverter()
        output_path = OUTPUT_DIR / f"{pdf_file.stem}_formatted.docx"
        
        converter.convert(
            str(pdf_file),
            str(output_path),
            preserve_format=True
        )
        
        result['processing_time'] = time.time() - start_time
        
        # 提取并对比
        generated_text = extract_docx_text(output_path)
        reference_text = extract_docx_text(ref_docx)
        result['text_similarity'] = calculate_similarity(generated_text, reference_text)
        
        generated_struct = extract_docx_structure(output_path)
        reference_struct = extract_docx_structure(ref_docx)
        
        result['generated_paragraphs'] = generated_struct['paragraph_count']
        result['reference_paragraphs'] = reference_struct['paragraph_count']
        result['generated_tables'] = generated_struct['table_count']
        result['reference_tables'] = reference_struct['table_count']
        result['generated_images'] = generated_struct['image_count']
        result['reference_images'] = reference_struct['image_count']
        
        # 结构相似度
        para_sim = min(generated_struct['paragraph_count'], reference_struct['paragraph_count']) / \
                   max(generated_struct['paragraph_count'], reference_struct['paragraph_count']) \
                   if max(generated_struct['paragraph_count'], reference_struct['paragraph_count']) > 0 else 1.0
        
        table_sim = min(generated_struct['table_count'], reference_struct['table_count']) / \
                    max(generated_struct['table_count'], reference_struct['table_count']) \
                    if max(generated_struct['table_count'], reference_struct['table_count']) > 0 else 1.0
        
        image_sim = min(generated_struct['image_count'], reference_struct['image_count']) / \
                    max(generated_struct['image_count'], reference_struct['image_count']) \
                    if max(generated_struct['image_count'], reference_struct['image_count']) > 0 else 1.0
        
        result['structure_similarity'] = (para_sim + table_sim + image_sim) / 3
        result['success'] = True
        
    except Exception as e:
        result['error'] = str(e)
    
    return result


def main():
    """主函数"""
    print("=" * 80)
    print("Phase 5 最终测试: 保留原格式模式下的PDF处理")
    print("=" * 80)
    
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    
    pdf_files = get_pdf_files()
    print(f"\n找到 {len(pdf_files)} 个PDF文件\n")
    
    results = []
    total_start = time.time()
    
    for i, pdf_file in enumerate(pdf_files, 1):
        print(f"[{i:2d}/{len(pdf_files)}] {pdf_file.name[:50]}...", end=" ")
        ref_docx = find_reference_docx(pdf_file)
        result = test_single_pdf(pdf_file, ref_docx)
        results.append(result)
        
        if result['success']:
            print(f"✅ 文本:{result['text_similarity']*100:5.1f}% 结构:{result['structure_similarity']*100:5.1f}% 时间:{result['processing_time']:5.1f}s")
        else:
            print(f"❌ {result['error']}")
    
    total_time = time.time() - total_start
    
    # 统计
    success_count = sum(1 for r in results if r['success'])
    text_sims = [r['text_similarity'] for r in results if r['success']]
    struct_sims = [r['structure_similarity'] for r in results if r['success']]
    times = [r['processing_time'] for r in results if r['success']]
    
    print("\n" + "=" * 80)
    print("测试结果汇总")
    print("=" * 80)
    print(f"总计: {len(results)} 个PDF")
    print(f"成功: {success_count}")
    print(f"失败: {len(results) - success_count}")
    print(f"总用时: {total_time:.1f}s")
    print(f"平均处理时间: {sum(times)/len(times):.1f}s" if times else "N/A")
    print()
    
    if text_sims:
        print(f"文本相似度: 平均 {sum(text_sims)/len(text_sims)*100:.1f}% | 最高 {max(text_sims)*100:.1f}% | 最低 {min(text_sims)*100:.1f}%")
    if struct_sims:
        print(f"结构相似度: 平均 {sum(struct_sims)/len(struct_sims)*100:.1f}% | 最高 {max(struct_sims)*100:.1f}% | 最低 {min(struct_sims)*100:.1f}%")
    
    # 保存报告
    report_path = OUTPUT_DIR / "phase5_final_report.json"
    with open(report_path, 'w', encoding='utf-8') as f:
        json.dump({
            'summary': {
                'total': len(results),
                'success': success_count,
                'failed': len(results) - success_count,
                'total_time': total_time,
                'avg_processing_time': sum(times)/len(times) if times else 0,
                'avg_text_similarity': sum(text_sims)/len(text_sims) if text_sims else 0,
                'avg_structure_similarity': sum(struct_sims)/len(struct_sims) if struct_sims else 0,
            },
            'results': results
        }, f, ensure_ascii=False, indent=2)
    
    print(f"\n报告已保存: {report_path}")
    
    if success_count == len(results):
        print("\n🎉 所有PDF处理成功！")
        return 0
    else:
        print(f"\n⚠️ {len(results) - success_count} 个PDF处理失败")
        return 1


if __name__ == "__main__":
    sys.exit(main())
