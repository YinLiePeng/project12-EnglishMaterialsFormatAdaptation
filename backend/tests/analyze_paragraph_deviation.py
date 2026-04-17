"""段落偏差分析脚本 - 分析4个偏差较大的PDF文件的段落计数差异原因"""

import sys
import os
import re
import json
import statistics
from pathlib import Path
from collections import defaultdict

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

BASE_DIR = os.path.dirname(os.path.dirname(__file__))
PDF_DIR = os.path.join(os.path.dirname(BASE_DIR), "测试用例", "原生PDF")
DOCX_DIR = os.path.join(os.path.dirname(BASE_DIR), "测试用例", "网上下载的英语资料样例")

TARGET_FILES = [
    "Unit 1 Art and artists单元话题练（语法填空+语法选择+阅读+完形+写作）（原卷版）.pdf",
    "Unit 1 Art and artists（单元测试·提升卷）英语新教材沪教版五四制八年级下册（原卷版）.pdf",
    "Unit3 Food 一轮复习单词及派生词语境记忆法导学案-2026届高三英语沪外版必修第二册[55909778].pdf",
    "Unit1 everyone is different 单元知识梳理&测试题（解析版） Unit1 everyone is different 单元知识梳理&测试题（学生版）.pdf",
]


def analyze_pdf_pages(pdf_path):
    """分析PDF每页的raw_lines数量和merge后的段落数量"""
    from app.services.pdf.parser import PDFParser
    from app.services.docx.parser import ElementType

    parser = PDFParser(pdf_path)
    try:
        page_count = parser.get_page_count()
        print(f"  PDF总页数: {page_count}")

        # 1) 检测header/footer regions
        hf_regions = parser._detect_header_footer_regions()
        table_regions = parser._collect_table_regions()

        # 2) 逐页提取 raw_lines 和 filtered_lines
        page_raw_counts = {}
        page_filtered_counts = {}
        page_hf_removed = {}
        page_raw_line_texts = {}
        page_filtered_line_texts = {}
        page_table_regions = {}

        for pn in range(page_count):
            page = parser.doc[pn]
            page_rect = page.rect
            columns = parser.detect_columns(pn)
            is_multi_col = len(columns) > 1

            raw_lines = parser._extract_raw_lines(pn, page_rect)
            page_raw_counts[pn] = len(raw_lines)
            page_raw_line_texts[pn] = [rl["text"][:80] for rl in raw_lines]

            # filter header/footer
            hf_for_page = hf_regions.get(pn, [])
            filtered = parser._filter_header_footer_lines(
                raw_lines, page_rect, hf_for_page
            )
            page_filtered_counts[pn] = len(filtered)
            page_hf_removed[pn] = len(raw_lines) - len(filtered)
            page_filtered_line_texts[pn] = [rl["text"][:80] for rl in filtered]

            # table regions on this page
            tr_this_page = [tr for tr in table_regions if tr["page_num"] == pn]
            page_table_regions[pn] = len(tr_this_page)

        # 3) 构建全量raw_lines（与convert_to_content_elements相同的流程）
        all_raw_lines = []
        for pn in range(page_count):
            page = parser.doc[pn]
            page_rect = page.rect
            columns = parser.detect_columns(pn)
            is_multi_col = len(columns) > 1

            raw_lines = parser._extract_raw_lines(pn, page_rect)
            hf_for_page = hf_regions.get(pn, [])
            filtered = parser._filter_header_footer_lines(
                raw_lines, page_rect, hf_for_page
            )

            if is_multi_col:
                col_lines = parser._assign_raw_lines_to_columns(filtered, columns)
                for col_idx in sorted(col_lines.keys()):
                    for rl in col_lines[col_idx]:
                        rl["column_idx"] = col_idx
                        all_raw_lines.append(rl)
            else:
                for rl in filtered:
                    rl["column_idx"] = 0
                    all_raw_lines.append(rl)

        all_raw_lines.sort(
            key=lambda x: (x["page_num"], x["column_idx"], x["y_position"])
        )

        total_raw = len(all_raw_lines)
        print(f"  全部raw_lines(过滤后): {total_raw}")

        # 4) 合并为段落
        paragraphs = parser._adaptive_merge_lines(all_raw_lines)
        print(f"  合并后段落数: {len(paragraphs)}")

        # 按页统计合并后段落数
        page_para_counts = defaultdict(int)
        for para in paragraphs:
            page_para_counts[para["page_num"]] += 1

        # 5) 最终elements
        elements = parser.convert_to_content_elements()
        para_elements = [e for e in elements if e.element_type == ElementType.PARAGRAPH]
        table_elements = [e for e in elements if e.element_type == ElementType.TABLE]
        image_elements = [e for e in elements if e.element_type == ElementType.IMAGE]
        print(
            f"  最终elements: {len(para_elements)}段落, {len(table_elements)}表格, {len(image_elements)}图片"
        )

        # 统计被table region过滤掉的段落数
        table_filtered_count = 0
        for para in paragraphs:
            para_spans = para.get("spans", [])
            x_pos = para_spans[0]["bbox"][0] if para_spans else 0
            x_end = para_spans[-1]["bbox"][2] if para_spans else 0
            pg = para["page_num"]
            y_pos = para.get("y_position", 0)
            if parser._is_in_table_region(pg, y_pos, table_regions, x_pos, x_end):
                table_filtered_count += 1
        print(f"  被table region过滤的段落数: {table_filtered_count}")

        return {
            "page_count": page_count,
            "total_raw": total_raw,
            "total_merged": len(paragraphs),
            "total_elements": len(para_elements),
            "table_elements": len(table_elements),
            "image_elements": len(image_elements),
            "table_filtered": table_filtered_count,
            "page_raw_counts": page_raw_counts,
            "page_filtered_counts": page_filtered_counts,
            "page_hf_removed": page_hf_removed,
            "page_para_counts": dict(page_para_counts),
            "page_table_regions": page_table_regions,
            "hf_regions": {
                k: [(r[0], r[1], r[2]) for r in v] for k, v in hf_regions.items()
            },
            "page_raw_line_texts": page_raw_line_texts,
            "page_filtered_line_texts": page_filtered_line_texts,
            "paragraphs": paragraphs,
            "columns_per_page": {},
        }
    finally:
        parser.close()


def analyze_docx_pages(docx_path):
    """分析标准DOCX的段落数"""
    from docx import Document

    doc = Document(docx_path)
    all_paras = doc.paragraphs
    text_paras = [p for p in all_paras if p.text.strip()]
    blank_paras = [p for p in all_paras if not p.text.strip()]

    tables = doc.tables
    table_cell_texts = []
    for t in tables:
        for row in t.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt:
                    table_cell_texts.append(txt)

    return {
        "total_paragraphs": len(all_paras),
        "text_paragraphs": len(text_paras),
        "blank_paragraphs": len(blank_paras),
        "tables": len(tables),
        "table_cell_texts_count": len(table_cell_texts),
        "text_sample": [p.text[:100] for p in text_paras[:15]],
    }


def detailed_page_comparison(pdf_path, docx_path, pdf_analysis):
    """详细对比差异最大的页面"""
    from app.services.pdf.parser import PDFParser
    from app.services.docx.parser import ElementType

    parser = PDFParser(pdf_path)
    try:
        # 获取最终elements按页分组
        elements = parser.convert_to_content_elements()

        page_elements = defaultdict(list)
        for e in elements:
            if e.element_type == ElementType.PARAGRAPH and e.paragraph:
                # 从runs推断页码（或者直接用原始数据）
                page_elements["all"].append(e.paragraph.text[:100])

        # 获取合并前后对比
        hf_regions = parser._detect_header_footer_regions()
        table_regions = parser._collect_table_regions()

        for pn in range(parser.get_page_count()):
            page = parser.doc[pn]
            page_rect = page.rect

            raw_lines = parser._extract_raw_lines(pn, page_rect)
            hf_for_page = hf_regions.get(pn, [])
            filtered = parser._filter_header_footer_lines(
                raw_lines, page_rect, hf_for_page
            )

            # 被hf过滤掉的行
            removed_texts = []
            for rl in raw_lines:
                is_removed = True
                for fl in filtered:
                    if (
                        fl["y_position"] == rl["y_position"]
                        and fl["text"] == rl["text"]
                    ):
                        is_removed = False
                        break
                if is_removed:
                    removed_texts.append(rl["text"][:80])

            if removed_texts:
                print(f"\n    --- 第{pn + 1}页 被HF过滤的行 ({len(removed_texts)}) ---")
                for t in removed_texts[:10]:
                    print(f"      [HF移除] {t}")
                if len(removed_texts) > 10:
                    print(f"      ... 还有{len(removed_texts) - 10}行")

        # 展示合并前后的段落对比
        print(f"\n  === 段落合并详情（前40个段落） ===")
        all_raw_lines = []
        for pn in range(parser.get_page_count()):
            page = parser.doc[pn]
            page_rect = page.rect
            columns = parser.detect_columns(pn)
            is_multi_col = len(columns) > 1

            raw_lines = parser._extract_raw_lines(pn, page_rect)
            hf_for_page = hf_regions.get(pn, [])
            filtered = parser._filter_header_footer_lines(
                raw_lines, page_rect, hf_for_page
            )

            if is_multi_col:
                col_lines = parser._assign_raw_lines_to_columns(filtered, columns)
                for col_idx in sorted(col_lines.keys()):
                    for rl in col_lines[col_idx]:
                        rl["column_idx"] = col_idx
                        all_raw_lines.append(rl)
            else:
                for rl in filtered:
                    rl["column_idx"] = 0
                    all_raw_lines.append(rl)

        all_raw_lines.sort(
            key=lambda x: (x["page_num"], x["column_idx"], x["y_position"])
        )

        paragraphs = parser._adaptive_merge_lines(all_raw_lines)

        for i, para in enumerate(paragraphs[:40]):
            text = para["text"]
            pg = para["page_num"] + 1
            spans = para["spans"]
            n_lines_merged = len([s for s in spans])
            # 计算合并了几条raw_line
            # 通过spans中有多少个原始line来估计
            if n_lines_merged > 2:
                print(f"    段{i}: P{pg} [{n_lines_merged} spans] {text[:100]}")
            else:
                print(f"    段{i}: P{pg} {text[:100]}")
        if len(paragraphs) > 40:
            print(f"    ... 共{len(paragraphs)}段，省略剩余")

    finally:
        parser.close()


def analyze_merge_quality(pdf_path):
    """分析合并质量：哪些本该合并没有合并，哪些不该合并却被合并了"""
    from app.services.pdf.parser import PDFParser

    parser = PDFParser(pdf_path)
    try:
        hf_regions = parser._detect_header_footer_regions()
        table_regions = parser._collect_table_regions()

        all_raw_lines = []
        for pn in range(parser.get_page_count()):
            page = parser.doc[pn]
            page_rect = page.rect
            columns = parser.detect_columns(pn)
            is_multi_col = len(columns) > 1

            raw_lines = parser._extract_raw_lines(pn, page_rect)
            hf_for_page = hf_regions.get(pn, [])
            filtered = parser._filter_header_footer_lines(
                raw_lines, page_rect, hf_for_page
            )

            if is_multi_col:
                col_lines = parser._assign_raw_lines_to_columns(filtered, columns)
                for col_idx in sorted(col_lines.keys()):
                    for rl in col_lines[col_idx]:
                        rl["column_idx"] = col_idx
                        all_raw_lines.append(rl)
            else:
                for rl in filtered:
                    rl["column_idx"] = 0
                    all_raw_lines.append(rl)

        all_raw_lines.sort(
            key=lambda x: (x["page_num"], x["column_idx"], x["y_position"])
        )

        # 计算主导间距
        line_gaps = parser._compute_dominant_gaps(all_raw_lines)

        print(f"\n  === 合并间距分析 ===")
        # 按页/栏显示主导间距
        for key in sorted(line_gaps.keys()):
            pn, col = key
            print(f"    Page{pn + 1} Col{col}: dominant_gap = {line_gaps[key]:.1f}")

        # 分析相邻行的gap分布
        gap_distribution = defaultdict(int)
        for i in range(1, len(all_raw_lines)):
            prev = all_raw_lines[i - 1]
            curr = all_raw_lines[i]
            if prev["page_num"] != curr["page_num"]:
                continue
            if prev.get("column_idx", 0) != curr.get("column_idx", 0):
                continue
            prev_fs = prev["spans"][0]["font_size"] if prev["spans"] else 12
            curr_fs = curr["spans"][0]["font_size"] if curr["spans"] else 12
            if abs(prev_fs - curr_fs) > 1.0:
                continue
            y_gap = curr["y_position"] - prev["y_end"]
            if y_gap > 0:
                bucket = round(y_gap, 0)
                gap_distribution[bucket] += 1

        print(f"\n  === Y间距分布 ===")
        for gap_val in sorted(gap_distribution.keys()):
            count = gap_distribution[gap_val]
            bar = "█" * min(count, 50)
            print(f"    gap={gap_val:6.0f}: {count:3d} {bar}")

        # 找出被合并的相邻行
        paragraphs = parser._adaptive_merge_lines(all_raw_lines)
        merged_groups = []
        current_start = 0
        for i, para in enumerate(paragraphs):
            n_spans = len(para["spans"])
            if n_spans > 1:
                # 检查这个段落包含了多少个原始raw_line
                # 通过y_position的跳变来判断
                prev_y = None
                line_count = 1
                for s in para["spans"]:
                    y = round(s["bbox"][1], 0)
                    if prev_y is not None and abs(y - prev_y) > 2:
                        line_count += 1
                    prev_y = y
                if line_count > 1:
                    merged_groups.append(
                        (i, line_count, para["text"][:80], para["page_num"] + 1)
                    )

        if merged_groups:
            print(f"\n  === 被合并的多行段落 ({len(merged_groups)}个) ===")
            for idx, lc, txt, pg in merged_groups[:30]:
                print(f"    段{idx}: P{pg} [{lc}行合并] {txt}")
            if len(merged_groups) > 30:
                print(f"    ... 还有{len(merged_groups) - 30}个")

    finally:
        parser.close()


def main():
    print("=" * 100)
    print("段落偏差分析报告")
    print("=" * 100)

    for pdf_name in TARGET_FILES:
        pdf_path = os.path.join(PDF_DIR, pdf_name)
        docx_name = pdf_name.replace(".pdf", ".docx")
        docx_path = os.path.join(DOCX_DIR, docx_name)

        print(f"\n{'=' * 100}")
        print(f"文件: {pdf_name}")
        print(f"{'=' * 100}")

        if not os.path.exists(pdf_path):
            print(f"  ERROR: PDF文件不存在: {pdf_path}")
            continue
        if not os.path.exists(docx_path):
            print(f"  ERROR: DOCX文件不存在: {docx_path}")
            continue

        # 1. 分析PDF
        print(f"\n  --- PDF分析 ---")
        pdf_info = analyze_pdf_pages(pdf_path)

        # 2. 分析标准DOCX
        print(f"\n  --- 标准DOCX分析 ---")
        docx_info = analyze_docx_pages(docx_path)
        print(f"  总段落: {docx_info['total_paragraphs']}")
        print(f"  文本段落: {docx_info['text_paragraphs']}")
        print(f"  空段落: {docx_info['blank_paragraphs']}")
        print(f"  表格: {docx_info['tables']}")

        # 3. 差异对比
        print(f"\n  --- 差异对比 ---")
        my_paras = pdf_info["total_elements"]
        std_paras = docx_info["text_paragraphs"]
        diff = my_paras - std_paras
        print(f"  PDF解析段落: {my_paras}")
        print(f"  标准DOCX段落: {std_paras}")
        print(f"  差异: {diff:+d}")

        # 逐页对比
        print(f"\n  --- 逐页对比 ---")
        print(
            f"  {'页码':>4s} | {'raw_lines':>9s} | {'filtered':>8s} | {'HF移除':>6s} | {'合并后段':>8s} | {'表格区':>6s}"
        )
        print(
            f"  {'-' * 4} | {'-' * 9} | {'-' * 8} | {'-' * 6} | {'-' * 8} | {'-' * 6}"
        )
        for pn in range(pdf_info["page_count"]):
            raw = pdf_info["page_raw_counts"].get(pn, 0)
            filt = pdf_info["page_filtered_counts"].get(pn, 0)
            hf = pdf_info["page_hf_removed"].get(pn, 0)
            merged = pdf_info["page_para_counts"].get(pn, 0)
            tables = pdf_info["page_table_regions"].get(pn, 0)
            print(
                f"  {pn + 1:4d} | {raw:9d} | {filt:8d} | {hf:6d} | {merged:8d} | {tables:6d}"
            )

        # 4. HF详情
        if pdf_info["hf_regions"]:
            print(f"\n  --- 页眉页脚区域详情 ---")
            for pn, regions in sorted(pdf_info["hf_regions"].items()):
                for y0, y1, rtype in regions:
                    print(f"    Page{pn + 1}: y={y0:.1f}~{y1:.1f} type={rtype}")

        # 5. 详细页面对比
        print(f"\n  --- 被HF过滤的行详情 ---")
        detailed_page_comparison(pdf_path, docx_path, pdf_info)

        # 6. 合并质量分析
        analyze_merge_quality(pdf_path)

        # 7. 标准DOCX前15段内容
        print(f"\n  --- 标准DOCX前30段内容 ---")
        for i, t in enumerate(docx_info["text_sample"][:30]):
            print(f"    [{i:3d}] {t}")
        if len(docx_info["text_sample"]) > 30:
            print(f"    ... 共{docx_info['text_paragraphs']}段")

    print(f"\n{'=' * 100}")
    print("分析完成")
    print(f"{'=' * 100}")


if __name__ == "__main__":
    main()
