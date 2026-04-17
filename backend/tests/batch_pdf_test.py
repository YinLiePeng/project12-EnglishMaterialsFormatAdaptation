"""批量PDF解析测试脚本 - 处理19个原生PDF并输出DOCX，用于效果评估"""

import sys
import os
import json
import traceback
from pathlib import Path
from datetime import datetime

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

BASE_DIR = os.path.dirname(os.path.dirname(__file__))
TEST_DATA_DIR = os.path.join(os.path.dirname(BASE_DIR), "测试用例")
NATIVE_PDF_DIR = os.path.join(TEST_DATA_DIR, "原生PDF")
STANDARD_DOCX_DIR = os.path.join(TEST_DATA_DIR, "网上下载的英语资料样例")
OUTPUT_DIR = os.path.join(BASE_DIR, "data", "pdf_test_output")


def process_single_pdf(pdf_path, output_dir):
    from app.services.pdf.parser import PDFParser
    from app.services.docx.parser import ElementType
    from app.services.docx import DocxGenerator
    from app.core.presets.styles import (
        get_preset_style,
        get_style_mapping,
        is_preserve_style,
    )

    filename = Path(pdf_path).stem
    parser = PDFParser(pdf_path)

    try:
        elements = parser.convert_to_content_elements()

        stats = {
            "total_elements": len(elements),
            "paragraphs": 0,
            "tables": 0,
            "images": 0,
            "blank_lines": 0,
        }
        for e in elements:
            if e.element_type == ElementType.PARAGRAPH:
                stats["paragraphs"] += 1
            elif e.element_type == ElementType.TABLE:
                stats["tables"] += 1
            elif e.element_type == ElementType.IMAGE:
                stats["images"] += 1
            elif e.element_type == ElementType.BLANK_LINE:
                stats["blank_lines"] += 1

        preserve = True
        style_mapping = {}
        style_keys = {}

        output_path = os.path.join(output_dir, f"{filename}_preserve.docx")
        generator = DocxGenerator()
        generator.generate_from_elements(
            elements, style_mapping, style_keys, preserve_format=preserve
        )
        generator.save(output_path)

        stats["output_path"] = output_path
        stats["output_size"] = (
            os.path.getsize(output_path) if os.path.exists(output_path) else 0
        )

        return True, stats, ""
    except Exception as e:
        return False, {}, traceback.format_exc()
    finally:
        parser.close()


def analyze_standard_docx(docx_path):
    from docx import Document

    doc = Document(docx_path)
    stats = {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "sections": len(doc.sections),
    }

    text_paras = [p for p in doc.paragraphs if p.text.strip()]
    stats["text_paragraphs"] = len(text_paras)
    stats["blank_paragraphs"] = len(doc.paragraphs) - len(text_paras)

    if doc.sections:
        section = doc.sections[0]
        stats["page_width_cm"] = (
            round(section.page_width / 360000, 2) if section.page_width else None
        )
        stats["page_height_cm"] = (
            round(section.page_height / 360000, 2) if section.page_height else None
        )
        stats["left_margin_cm"] = (
            round(section.left_margin / 360000, 2) if section.left_margin else None
        )
        stats["right_margin_cm"] = (
            round(section.right_margin / 360000, 2) if section.right_margin else None
        )
        stats["top_margin_cm"] = (
            round(section.top_margin / 360000, 2) if section.top_margin else None
        )
        stats["bottom_margin_cm"] = (
            round(section.bottom_margin / 360000, 2) if section.bottom_margin else None
        )

    font_sizes = set()
    font_names = set()
    for p in text_paras[:50]:
        for run in p.runs:
            if run.font.size:
                font_sizes.add(round(run.font.size.pt, 1))
            if run.font.name:
                font_names.add(run.font.name)

    stats["font_sizes"] = sorted(font_sizes)
    stats["font_names"] = sorted(font_names)

    if text_paras:
        sample_texts = []
        for p in text_paras[:10]:
            align = str(p.alignment) if p.alignment else "None"
            runs_info = []
            for r in p.runs[:3]:
                runs_info.append(
                    {
                        "text": r.text[:50],
                        "font": r.font.name,
                        "size": round(r.font.size.pt, 1) if r.font.size else None,
                        "bold": r.font.bold,
                    }
                )
            sample_texts.append(
                {
                    "text": p.text[:80],
                    "alignment": align,
                    "runs": runs_info,
                }
            )
        stats["sample_paragraphs"] = sample_texts

    return stats


def extract_pdf_structure(pdf_path):
    from app.services.pdf.parser import PDFParser

    parser = PDFParser(pdf_path)
    try:
        elements = parser.convert_to_content_elements()
        from app.services.docx.parser import ElementType

        result = []
        for e in elements[:30]:
            if e.element_type == ElementType.PARAGRAPH and e.paragraph:
                p = e.paragraph
                runs_brief = []
                for r in p.runs[:5]:
                    runs_brief.append(
                        {
                            "text": r.text[:40],
                            "font": r.font_name,
                            "size": r.font_size,
                            "bold": r.bold,
                        }
                    )
                result.append(
                    {
                        "type": "paragraph",
                        "text": p.text[:80],
                        "font": {
                            "name": p.font.name,
                            "size": p.font.size,
                            "bold": p.font.bold,
                        },
                        "alignment": p.format.alignment,
                        "runs_count": len(p.runs),
                        "runs_sample": runs_brief,
                    }
                )
            elif e.element_type == ElementType.TABLE and e.table_cells:
                rows = len(e.table_cells)
                cols = max((len(r) for r in e.table_cells), default=0)
                result.append({"type": "table", "rows": rows, "cols": cols})
            elif e.element_type == ElementType.IMAGE:
                result.append(
                    {"type": "image", "size": len(e.image_data) if e.image_data else 0}
                )

        return result
    finally:
        parser.close()


if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    pdf_files = sorted([f for f in os.listdir(NATIVE_PDF_DIR) if f.endswith(".pdf")])

    print(f"=" * 80)
    print(f"PDF批量测试 - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"共 {len(pdf_files)} 个PDF文件")
    print(f"输出目录: {OUTPUT_DIR}")
    print(f"=" * 80)

    all_results = {}

    for i, pdf_name in enumerate(pdf_files):
        pdf_path = os.path.join(NATIVE_PDF_DIR, pdf_name)
        docx_name = pdf_name.replace(".pdf", ".docx")
        standard_path = os.path.join(STANDARD_DOCX_DIR, docx_name)

        print(f"\n[{i + 1}/{len(pdf_files)}] {pdf_name[:60]}...")

        success, stats, error = process_single_pdf(pdf_path, OUTPUT_DIR)

        result = {"pdf": pdf_name, "success": success, "stats": stats, "error": error}

        if os.path.exists(standard_path):
            try:
                std_stats = analyze_standard_docx(standard_path)
                result["standard"] = std_stats
            except Exception as e:
                result["standard_error"] = str(e)

        try:
            pdf_struct = extract_pdf_structure(pdf_path)
            result["pdf_structure"] = pdf_struct
        except Exception as e:
            result["pdf_structure_error"] = str(e)

        all_results[pdf_name] = result

        if success:
            print(
                f"  OK - {stats['paragraphs']}段/{stats['tables']}表/{stats['images']}图 -> {stats['output_size']}bytes"
            )
            if "standard" in result:
                std = result["standard"]
                print(
                    f"  标准: {std['paragraphs']}段/{std['tables']}表 字体:{std.get('font_sizes', [])[:5]}"
                )
        else:
            print(f"  FAILED - {error[:100]}")

    report_path = os.path.join(OUTPUT_DIR, "batch_report.json")
    with open(report_path, "w", encoding="utf-8") as f:
        json.dump(all_results, f, ensure_ascii=False, indent=2, default=str)

    print(f"\n{'=' * 80}")
    print(f"报告已保存: {report_path}")

    print(f"\n\n{'=' * 80}")
    print("对比摘要:")
    print(f"{'=' * 80}")
    for pdf_name, result in all_results.items():
        if not result["success"]:
            continue
        std = result.get("standard", {})
        my = result["stats"]
        para_diff = my["paragraphs"] - std.get("text_paragraphs", 0)
        table_diff = my["tables"] - std.get("tables", 0)
        img_diff = my["images"] - sum(
            1 for p in std.get("sample_paragraphs", []) if "图片" in p.get("text", "")
        )
        print(
            f"  {pdf_name[:50]:50s} | 段落:{my['paragraphs']:3d}(标{std.get('text_paragraphs', '?'):>3s}) 表格:{my['tables']:2d}(标{std.get('tables', '?'):>2s}) 图片:{my['images']:2d}"
        )
